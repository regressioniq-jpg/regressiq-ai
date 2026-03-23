
import json
from io import BytesIO
import streamlit as st
from docx import Document

from pipeline import graph
from vector_store import ingest_tests
from visualizer import risk_heatmap, dependency_graph


# ---------------- DOCX GENERATOR ----------------
def generate_docx_report(data):
    doc = Document()

    # Title
    doc.add_heading('RegressIQ AI - Regression Report', 0)

    # Change Request
    doc.add_heading('Change Request', level=1)
    doc.add_paragraph(data["cr_text"])

    # Change Analysis
    doc.add_heading('Change Analysis', level=1)
    ca = data["change_analysis"]
    doc.add_paragraph(f"Changed Modules: {', '.join(ca['changed_modules'])}")
    doc.add_paragraph(f"Change Type: {ca['change_type']}")
    doc.add_paragraph(f"Summary: {ca['summary']}")

    # Impacted Modules
    doc.add_heading('Impacted Modules', level=1)
    doc.add_paragraph(", ".join(data["impacted_modules"]))

    # Risk
    doc.add_heading('Risk Assessment', level=1)
    doc.add_paragraph(f"Overall Risk: {data['risk_assessment']['overall_risk']}")

    # Test Plan
    doc.add_heading('Recommended Test Cases', level=1)
    for test in data["test_plan"]:
        doc.add_paragraph(
            f"{test['test_id']} - {test['title']} ({test['module']})",
            style='List Bullet'
        )

    # Coverage
    doc.add_heading('Coverage', level=1)
    cov = data["coverage"]
    doc.add_paragraph(f"Coverage: {cov['coverage_percent']}%")

    # Strategy
    doc.add_heading('Regression Strategy', level=1)
    doc.add_paragraph(data["strategy"])

    return doc


# ---------------- UI ----------------
st.set_page_config(page_title="RegressIQ AI", layout="wide")

st.title("🚀 RegressIQ AI")
st.caption("AI-powered regression impact analysis with LangGraph + LLM + Vector DB")

# ---------------- SIDEBAR ----------------
with st.sidebar:
    st.header("⚙️ Setup")

    if st.button("📥 Index Test Cases"):
        added = ingest_tests()
        st.success(f"Indexed {added} new test case(s).")

    st.markdown("### 🎯 Demo Scenarios")

    demo_1 = "Payment service retry logic updated and order confirmation callback modified to prevent duplicate payments."
    demo_2 = "Fraud detection threshold updated and notification service rules changed for suspicious transactions."
    demo_3 = "Inventory service reservation logic updated during checkout affecting order and payment flow."

    if st.button("💳 Payment Issue"):
        st.session_state["cr_text"] = demo_1

    if st.button("🛡️ Fraud Update"):
        st.session_state["cr_text"] = demo_2

    if st.button("📦 Inventory Issue"):
        st.session_state["cr_text"] = demo_3


# ---------------- INPUT ----------------
cr_text = st.text_area(
    "📝 Enter Change Request",
    value=st.session_state.get("cr_text", ""),
    height=150,
    placeholder="Paste the change request here..."
)

# ---------------- RUN ----------------
if st.button("🚀 Run Analysis", type="primary"):
    if not cr_text.strip():
        st.warning("Please enter a change request.")
    else:
        with st.spinner("🤖 AI analyzing change request..."):
            result = graph.invoke({"cr_text": cr_text})

        # ---------------- SUMMARY ----------------
        st.markdown("## 📊 Quick Summary")

        colA, colB, colC = st.columns(3)

        with colA:
            st.success(f"✅ Modules Impacted\n\n{len(result['impacted_modules'])}")

        with colB:
            st.warning(f"⚠️ Overall Risk\n\n{result['risk_assessment']['overall_risk']}")

        with colC:
            st.info(f"🧪 Tests Generated\n\n{len(result['test_plan'])}")

        st.divider()

        # ---------------- CHANGE + IMPACT ----------------
        col1, col2 = st.columns(2)

        with col1:
            st.subheader("🔧 Changed Modules")
            st.write(result["change_analysis"]["changed_modules"])

        with col2:
            st.subheader("🌐 Impacted Modules")
            st.write(result["impacted_modules"])

        st.divider()

        # ---------------- VISUALS ----------------
        col1, col2 = st.columns(2)

        with col1:
            st.subheader("📊 Risk Heatmap")
            st.plotly_chart(
                risk_heatmap(result["risk_assessment"]),
                width="stretch"
            )

        with col2:
            st.subheader("🕸️ Dependency Graph")
            st.plotly_chart(
                dependency_graph(
                    result["graph"],
                    result["impacted_modules"],
                    result["risk_assessment"]
                ),
                width="stretch"
            )

        st.divider()

        # ---------------- TESTS ----------------
        st.subheader("🧪 Recommended Regression Tests")
        st.dataframe(result["test_plan"], width="stretch")

        # ---------------- COVERAGE ----------------
        st.subheader("📈 Coverage")
        st.json(result["coverage"])

        # ---------------- STRATEGY ----------------
        st.subheader("🧠 Regression Strategy")
        st.markdown(result["strategy"])

        st.divider()

        # ---------------- DOWNLOAD DOCX ----------------
        export_result = {
            "cr_text": result["cr_text"],
            "change_analysis": result["change_analysis"],
            "impacted_modules": result["impacted_modules"],
            "risk_assessment": result["risk_assessment"],
            "test_plan": result["test_plan"],
            "coverage": result["coverage"],
            "strategy": result["strategy"]
        }

        doc = generate_docx_report(export_result)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            "⬇️ Download Report (DOCX)",
            data=buffer,
            file_name="regressiq_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

